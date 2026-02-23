import streamlit as st
from transformers import pipeline
from streamlit_option_menu import option_menu
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import A4
import sqlite3
import hashlib

st.set_page_config(page_title="AI Emotion Analyzer Pro", layout="wide")

# =====================================================
# 🔐 AUTHENTICATION SYSTEM (Login / Register / Reset)
# =====================================================

conn = sqlite3.connect("users.db", check_same_thread=False)
c = conn.cursor()

c.execute("""
CREATE TABLE IF NOT EXISTS users (
    username TEXT PRIMARY KEY,
    password TEXT
)
""")
conn.commit()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

def auth_page():
    option = st.selectbox("Select Option", ["Login", "Register", "Forgot Password"])

    # -------- LOGIN --------
    if option == "Login":
        st.title("🔐 Login")

        username = st.text_input("Username")
        password = st.text_input("Password", type="password")

        if st.button("Login"):
            c.execute("SELECT * FROM users WHERE username=? AND password=?",
                      (username, hash_password(password)))
            user = c.fetchone()

            if user:
                st.session_state.authenticated = True
                st.success("Login Successful!")
                st.rerun()
            else:
                st.error("Invalid Username or Password")

    # -------- REGISTER --------
    elif option == "Register":
        st.title("📝 Create Account")

        new_user = st.text_input("New Username")
        new_pass = st.text_input("New Password", type="password")

        if st.button("Register"):
            if new_user and new_pass:
                try:
                    c.execute("INSERT INTO users VALUES (?, ?)",
                              (new_user, hash_password(new_pass)))
                    conn.commit()
                    st.success("Account created successfully! Please login.")
                except:
                    st.error("Username already exists.")
            else:
                st.warning("Please fill all fields.")

    # -------- FORGOT PASSWORD --------
    elif option == "Forgot Password":
        st.title("🔄 Reset Password")

        user_name = st.text_input("Username")
        new_password = st.text_input("New Password", type="password")

        if st.button("Reset Password"):
            c.execute("SELECT * FROM users WHERE username=?", (user_name,))
            user = c.fetchone()

            if user:
                c.execute("UPDATE users SET password=? WHERE username=?",
                          (hash_password(new_password), user_name))
                conn.commit()
                st.success("Password updated! Please login.")
            else:
                st.error("Username not found.")

def logout():
    st.session_state.authenticated = False
    st.rerun()

# =====================================================
# 📄 REPORT GENERATION
# =====================================================

def create_docx(text, top_emotion, confidence, chart_buf):
    doc = Document()
    doc.add_heading("Emotion Analysis Report", 0)
    doc.add_paragraph(f"Text: {text}")
    doc.add_paragraph(f"Predicted Emotion: {top_emotion}")
    doc.add_paragraph(f"Confidence: {confidence:.2f}%")
    doc.add_picture(chart_buf, width=Inches(4))
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_pdf(text, top_emotion, confidence, chart_buf):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("<b>Emotion Analysis Report</b>", styles["Title"]))
    elements.append(Spacer(1, 0.3 * inch))
    elements.append(Paragraph(f"<b>Text:</b> {text}", styles["Normal"]))
    elements.append(Spacer(1, 0.2 * inch))
    elements.append(Paragraph(f"<b>Predicted Emotion:</b> {top_emotion}", styles["Normal"]))
    elements.append(Spacer(1, 0.2 * inch))
    elements.append(Paragraph(f"<b>Confidence:</b> {confidence:.2f}%", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    chart_buf.seek(0)
    img = RLImage(chart_buf, width=5 * inch, height=3 * inch)
    elements.append(img)

    doc.build(elements)
    buffer.seek(0)
    return buffer

# =====================================================
# 🤖 LOAD MODEL
# =====================================================

@st.cache_resource
def load_model():
    return pipeline(
        "text-classification",
        model="bhadresh-savani/distilbert-base-uncased-emotion",
        return_all_scores=True
    )

classifier = load_model()

# =====================================================
# 📜 STORE HISTORY
# =====================================================

if "history" not in st.session_state:
    st.session_state.history = []

# =====================================================
# 🔐 CHECK LOGIN
# =====================================================

if not st.session_state.authenticated:
    auth_page()
    st.stop()

# =====================================================
# 🧭 NAVIGATION
# =====================================================

selected = option_menu(
    menu_title=None,
    options=["Dashboard", "Analyze", "History", "Compare", "Statistics", "About"],
    icons=["speedometer", "emoji-smile", "clock-history", "bar-chart", "graph-up", "info-circle"],
    orientation="horizontal",
)

st.write("---")
if st.button("🚪 Logout"):
    logout()

# =====================================================
# 📊 DASHBOARD
# =====================================================

if selected == "Dashboard":
    st.title("🤖 Advanced AI Emotion Analyzer")
    st.write("Understand the emotions behind any text instantly.")
    st.write("---")

    col1, col2, col3 = st.columns(3)
    col1.metric("Speed", "Instant Analysis")
    col2.metric("Reports", "PDF & DOCX Available")
    col3.metric("Interface", "Simple & Professional")

    st.write("---")
    st.write("""
    🔹 Analyze emotional tone in seconds  
    🔹 View confidence percentage  
    🔹 See graphical emotion distribution  
    🔹 Download professional reports  
    """)

# =====================================================
# 💬 ANALYZE
# =====================================================

elif selected == "Analyze":
    st.header("💬 Enter Text for Emotion Analysis")
    user_input = st.text_area("Type your text here", height=150)

    if st.button("Analyze Emotion"):
        if user_input.strip():
            results = classifier(user_input, return_all_scores=True)

            if isinstance(results[0], list):
                results = results[0]

            results = sorted(results, key=lambda x: x["score"], reverse=True)

            top_emotion = results[0]["label"]
            confidence = results[0]["score"] * 100

            emoji_map = {
                "joy": "😊",
                "sadness": "😢",
                "anger": "😠",
                "fear": "😨",
                "surprise": "😲",
                "love": "❤️"
            }

            st.success(f"{top_emotion.upper()} {emoji_map.get(top_emotion,'')}")
            st.write(f"Confidence: {confidence:.2f}%")

            st.session_state.history.append({
                "text": user_input,
                "emotion": top_emotion,
                "confidence": confidence
            })

            labels = [r["label"] for r in results]
            scores = [r["score"] * 100 for r in results]

            fig, ax = plt.subplots(figsize=(8, 4))
            ax.bar(labels, scores)
            ax.set_xlabel("Emotions")
            ax.set_ylabel("Confidence (%)")
            ax.set_title("Emotion Probability Distribution")
            st.pyplot(fig)

            buf = io.BytesIO()
            fig.savefig(buf, format="png")
            buf.seek(0)

            st.download_button("📄 Download DOCX",
                               create_docx(user_input, top_emotion, confidence, buf),
                               "emotion_report.docx")

            st.download_button("📑 Download PDF",
                               create_pdf(user_input, top_emotion, confidence, buf),
                               "emotion_report.pdf")

# =====================================================
# 📜 HISTORY
# =====================================================

elif selected == "History":
    st.title("📜 Analysis History")

    if st.session_state.history:

        if st.button("🗑 Clear History"):
            st.session_state.history = []
            st.rerun()

        for item in reversed(st.session_state.history):
            st.write(f"**Text:** {item['text']}")
            st.write(f"Emotion: {item['emotion']} | Confidence: {item['confidence']:.2f}%")
            st.write("---")
    else:
        st.info("No analysis history yet.")

# =====================================================
# 📊 COMPARE
# =====================================================

elif selected == "Compare":
    st.title("📊 Compare Two Texts")

    text1 = st.text_area("Enter First Text")
    text2 = st.text_area("Enter Second Text")

    if st.button("Compare Emotions"):
        if text1.strip() and text2.strip():
            result1 = classifier(text1)[0]
            result2 = classifier(text2)[0]

            st.success(f"Text 1 → {result1['label'].upper()} ({result1['score']*100:.2f}%)")
            st.success(f"Text 2 → {result2['label'].upper()} ({result2['score']*100:.2f}%)")
        else:
            st.warning("Please enter both texts.")

# =====================================================
# 📈 STATISTICS
# =====================================================

elif selected == "Statistics":
    st.title("📈 Emotion Statistics")

    if st.session_state.history:
        emotion_counts = {}
        for item in st.session_state.history:
            emotion = item["emotion"]
            emotion_counts[emotion] = emotion_counts.get(emotion, 0) + 1

        fig, ax = plt.subplots()
        ax.bar(emotion_counts.keys(), emotion_counts.values())
        ax.set_xlabel("Emotions")
        ax.set_ylabel("Frequency")
        ax.set_title("Emotion Frequency")
        st.pyplot(fig)
    else:
        st.info("No data available.")

# =====================================================
# ℹ️ ABOUT
# =====================================================

elif selected == "About":
    st.title("About This Application")

    st.write("""
    This application helps you detect emotional tone from text.

    🔹 Detect dominant emotion  
    🔹 View confidence scores  
    🔹 Visual emotion graphs  
    🔹 Download reports  
    🔹 Compare multiple texts  
    🔹 Track emotion history  

    Designed to be simple, fast, and professional.
    """)